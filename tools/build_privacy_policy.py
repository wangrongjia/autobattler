from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path(r"C:\Users\admin\Downloads\privacy-policy-template.docx")
OUTPUT = Path(r"D:\github\autobattler\TapTapAssets\战三国·弈定九州-隐私政策-待补全.docx")


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_para(doc, text="", bold_prefix=None, align=None, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=14, bold=True, color=(142, 62, 42))
    return p


def create_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet")
    level.append(fmt)
    text_el = OxmlElement("w:lvlText")
    text_el.set(qn("w:val"), "•")
    level.append(text_el)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    level.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, text, num_id):
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_el)
    ppr.append(num_pr)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    for r in p.runs:
        set_run_font(r)
    if not p.runs:
        r = p.add_run(text)
        set_run_font(r)
    else:
        p.runs[0].text = text
    return p


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r, size=9.5, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_table(table, widths):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "D8C8BB")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            margin = tc_pr.find(qn("w:tcMar"))
            if margin is None:
                margin = OxmlElement("w:tcMar")
                tc_pr.append(margin)
            for side in ("top", "left", "bottom", "right"):
                side_el = margin.find(qn(f"w:{side}"))
                if side_el is None:
                    side_el = OxmlElement(f"w:{side}")
                    margin.append(side_el)
                side_el.set(qn("w:w"), "100")
                side_el.set(qn("w:type"), "dxa")


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)

    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    bullet_num_id = create_bullet_numbering(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("《战三国·弈定九州》隐私政策")
    set_run_font(run, size=20, bold=True, color=(142, 62, 42))

    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.paragraph_format.space_after = Pt(14)
    run = status.add_run("待补全运营者信息版")
    set_run_font(run, size=11, bold=True, color=(180, 105, 35))

    add_para(doc, "更新日期：2026年8月27日", align=WD_ALIGN_PARAGRAPH.RIGHT, after=2)
    add_para(doc, "生效日期：2026年8月27日", align=WD_ALIGN_PARAGRAPH.RIGHT, after=12)

    p = add_para(
        doc,
        "待补充后再公开：运营者法定名称（个人开发者填写真实姓名）、联系地址、联系邮箱。请勿仅填写工作室昵称。",
        after=12,
    )
    for r in p.runs:
        set_run_font(r, bold=True, color=(156, 76, 0))
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.6)

    add_para(
        doc,
        "【请填写运营者法定名称】（品牌名称：序言工作室，以下简称“我们”）是游戏《战三国·弈定九州》的运营者。我们重视用户隐私，并依据实际功能说明本游戏如何处理信息。",
    )
    add_para(
        doc,
        "本游戏目前为单机游戏，不提供账号注册、实名认证、充值、广告投放、社交聊天或 AI 合成服务。除玩家主动使用“GitHub 私有 Gist 手动同步”功能外，游戏不会将本地存档或个人信息上传至我们的服务器。",
    )

    add_heading(doc, "一、我们如何收集和使用信息")
    add_para(doc, "1. 本地游戏数据", bold_prefix="1. 本地游戏数据")
    add_para(
        doc,
        "为提供保存进度、读取进度、设置、天赋与符文成长、阵容和平衡实验室等功能，游戏会在您的设备本地保存游戏进度、设置项和您主动创建的配置数据。这些数据默认仅保存在设备本地，我们不会自动上传。",
    )
    add_para(doc, "2. 可选的 GitHub 私有 Gist 同步", bold_prefix="2. 可选的 GitHub 私有 Gist 同步")
    add_para(
        doc,
        "平衡实验室提供可选的 GitHub 私有 Gist 手动同步功能。仅当您主动填写 GitHub Classic Token、Gist ID 并点击上传或下载时，游戏才会连接 GitHub API。Token 与 Gist ID 保存在您的设备本地；阵容数值覆盖文件会按照您的操作直接传输至您自己的 GitHub 私有 Gist。该功能由 GitHub 提供，我们不会接收或保存您的 GitHub Token、Gist ID 或同步内容。",
    )
    add_para(doc, "3. 我们不收集的信息", bold_prefix="3. 我们不收集的信息")
    for item in [
        "手机号、登录密码、真实姓名、身份证件号码或人脸信息；",
        "设备唯一标识符、通讯录、短信、通话记录、精确位置、相册、相机或麦克风内容；",
        "支付信息、广告标识符或用于个性化推荐的行为画像；",
        "由我们运营的服务器日志、统计分析或崩溃遥测数据。",
    ]:
        add_bullet(doc, item, bullet_num_id)

    add_heading(doc, "二、设备权限与网络访问")
    add_para(doc, "本游戏 APK 仅声明网络访问能力。该能力主要用于您主动选择的 GitHub 私有 Gist 同步；不使用该功能时，核心单机玩法、存档和设置无需联网。Android 的网络访问能力通常不会弹出运行时授权窗口。")
    table = doc.add_table(rows=2, cols=4)
    configure_table(table, [1700, 3000, 2500, 2160])
    headers = ["能力/权限", "使用目的", "是否主动触发", "能否避免"]
    values = ["网络访问", "可选的 GitHub 私有 Gist 上传与下载", "是，仅在用户点击同步时", "可以；不使用同步功能即可"]
    for i, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], text, bold=True, color=(255, 255, 255))
        shade_cell(table.rows[0].cells[i], "8E3E2A")
    for i, text in enumerate(values):
        set_cell_text(table.rows[1].cells[i], text)
        shade_cell(table.rows[1].cells[i], "F7EFE8")

    add_heading(doc, "三、信息的保存与删除")
    add_para(doc, "本地游戏数据保存在您的设备应用数据目录中，保存期限由您决定。您可通过游戏内的新游戏/重置功能、Android 系统的“清除存储数据”功能或卸载游戏删除本地数据。")
    add_para(doc, "如您使用 GitHub 私有 Gist 同步，相关数据的保存期限、删除方式和账号管理由您在 GitHub 中控制。您可以在 GitHub 删除对应 Gist、撤销 Token 或删除 GitHub 账号。")

    add_heading(doc, "四、Cookies 和同类技术")
    add_para(doc, "本游戏不使用 Cookies、网页跟踪像素或广告跟踪技术。")

    add_heading(doc, "五、共享、转让与公开披露")
    add_para(doc, "我们不会出售、出租、共享、转让或公开披露您的个人信息。只有在法律法规明确要求或有权机关依法提出要求时，我们才可能依法履行相应义务。可选 GitHub 同步是您与 GitHub 之间的直接传输，不属于我们向第三方共享您的信息。")

    add_heading(doc, "六、第三方服务")
    add_para(doc, "本游戏不接入广告、统计、支付或账号登录 SDK。可选同步功能使用 GitHub API。GitHub 对信息的处理适用其隐私声明：")
    add_para(doc, "GitHub Privacy Statement：https://docs.github.com/zh/site-policy/privacy-policies/github-general-privacy-statement")
    add_para(doc, "玩家通过 TapTap 下载或浏览游戏页面时，TapTap 可能按照其自身隐私政策处理相关信息；该处理由 TapTap 独立负责。")

    add_heading(doc, "七、我们如何保护信息")
    add_para(doc, "我们遵循数据最小化原则，尽可能让游戏数据仅保存在您的设备本地。请您妥善保管自行填写的 GitHub Token，不要向他人泄露，并建议只授予完成同步所需的最小权限。任何互联网传输都无法保证绝对安全，请您在使用可选同步功能前充分了解相关风险。")

    add_heading(doc, "八、您的权利")
    for item in [
        "访问和更正：您可以在游戏内查看并修改设置、进度和自定义配置；",
        "删除：您可以清除本地应用数据或卸载游戏；使用 GitHub 同步时，可在 GitHub 删除对应 Gist；",
        "撤回同意：不使用 GitHub 同步功能，或删除本地保存的 Token/Gist ID 并撤销 GitHub Token；",
        "投诉与咨询：您可以通过本政策末尾的联系方式联系我们。",
    ]:
        add_bullet(doc, item, bullet_num_id)

    add_heading(doc, "九、未成年人保护")
    add_para(doc, "本游戏适龄分级为 12 周岁以上。未成年人应在监护人指导下阅读本政策并使用游戏。我们不会主动收集未成年人的身份信息；如监护人发现未成年人向我们提供了个人信息，请通过本政策载明的方式联系我们，我们将在核实后依法处理。")

    add_heading(doc, "十、本政策的更新")
    add_para(doc, "游戏功能、信息处理方式或法律法规发生变化时，我们可能更新本政策。重大变更将通过游戏页面、更新说明或其他显著方式告知。更新后的政策将在公开页面标明更新日期和生效日期。")

    add_heading(doc, "十一、如何联系我们")
    contact_rows = [
        ("运营者法定名称", "【请填写；个人开发者填写真实姓名】"),
        ("品牌名称", "序言工作室"),
        ("联系地址", "【请填写可接收联系的地址】"),
        ("联系邮箱", "【请填写长期有效的客服邮箱】"),
    ]
    contact = doc.add_table(rows=len(contact_rows), cols=2)
    configure_table(contact, [2600, 6760])
    for row, (label, value) in zip(contact.rows, contact_rows):
        set_cell_text(row.cells[0], label, bold=True)
        shade_cell(row.cells[0], "F7EFE8")
        set_cell_text(row.cells[1], value)
    add_para(doc, "一般情况下，我们将在十五个工作日内回复您的请求。")

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.text = "《战三国·弈定九州》隐私政策 · 待补全运营者信息后公开"
    for r in footer_p.runs:
        set_run_font(r, size=8.5, color=(120, 120, 120))

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
